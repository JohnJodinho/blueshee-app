from typing import Iterable
from texts import SHAPE, SOUTH_WEST, MISCO, SYSTEM_INSTRUCTION
import io
import time
import pandas as pd
import json
# For Excel Files:
from openpyxl import Workbook  # Create and save new Excel workbooks
from openpyxl.utils import get_column_letter  # Adjust column widths by letter reference
from openpyxl.styles import Font, Alignment, PatternFill  # Styling and formatting cells
from openpyxl.worksheet.worksheet import Worksheet  # Handling worksheets specifically
from openpyxl.drawing.image import Image  # Insert images into Excel files (if needed)
from docx import Document
from docx.shared import Pt
# For General File Operations:
import os  # Handle file paths and directory operations
from os import path  # Check if files or directories exist, get file paths
# For alternative to excel:
import csv  # If generating CSV files as an alternative to Excel
# Utulity Libraries:
import datetime  # For adding timestamps to file names, if needed
import re  # Use regular expressions for text processing
import vertexai
from vertexai.preview.generative_models import (
    GenerationResponse,
    GenerativeModel,
    HarmBlockThreshold,
    HarmCategory,
    Part, 
    ChatSession
)


# Configuration
PROJECT_ID = "my-bluesheet-project-440016"  # Replace with your project ID
LOCATION = "us-central1"  # Replace with your location
MODEL_NAME = "gemini-1.5-pro-002"  # Replace with model name
BLOCK_LEVEL = HarmBlockThreshold.BLOCK_ONLY_HIGH
vertexai.init(project=PROJECT_ID, location=LOCATION)





# Load the document (Word, PDF, or Excel)
def load_document(file_path):
    file_extension = file_path.split('.')[-1].lower()
    
    mime_type = None
    if file_extension == "pdf":
        mime_type = "application/pdf"
    elif file_extension in [".csv", ".txt"]:
        mime_type = "text/plain"
        
    else:
        print("Unsupported file type. Please provide a PDF, Word, or Excel document.")
    
    # Load the file
    with open(file_path, "rb") as fp:
        document = Part.from_data(data=fp.read(), mime_type=mime_type)
    
    return document

def is_valid_json(response_text):
    """
    Check if response_text is valid JSON.
    
    Strips any JSON code block delimiters (e.g., ```json) and 
    then attempts to parse it to confirm validity.
    """
    # Remove JSON code block delimiters if they exist
    json_text = re.sub(r'```json|```', '', response_text).strip()

    try:
        # Try loading the text as JSON
        json.loads(json_text)
        return True
    except ValueError:
        # If parsing fails, it's not valid JSON
        return False
def save_to_excel(json_data):
    """Save JSON data to Excel using openpyxl."""
    from openpyxl import Workbook

    # Create a new workbook
    wb = Workbook()
    document_name = json_data.get("document_name", "Bluesheet Draft.xlsx")
    
    # Process each worksheet in the JSON data
    for sheet_name, rows in json_data.items():
        if sheet_name == "document_name":
            continue  # Skip the document name key
        ws = wb.create_sheet(title=sheet_name)
        
        # Assume first row in the first list entry has the headers
        headers = rows[0].keys()
        ws.append(list(headers))
        
        # Add each row from the JSON data
        for row_data in rows:
            ws.append(list(row_data.values()))
    
    # Save the workbook with the specified document name
    wb.save(f'{document_name}')

def json_to_docx(json_string: str):
    # Parse JSON string into a Python dictionary
    data = re.sub(r'```json|```', '', json_string).strip()
    data = json.loads(json_string)
    
    # Retrieve filename from the first dictionary
    file_name = data[0].get("fileName", "output.docx")
    
    # Initialize a Word document
    document = Document()
    
    # Iterate through JSON data (skip the first dictionary as it only contains the filename)
    for item in data[1:]:
        content_type = item.get("type")
        text = item.get("text", "")
        font_size = item.get("fontSize", 12)
        level = item.get("level", 1)  # For headings, this defines heading level

        if content_type == "heading":
            # Add a heading with the specified level
            heading = document.add_heading(text, level=level)
            # Adjust font size of heading if specified
            for run in heading.runs:
                run.font.size = Pt(font_size)
        
        elif content_type == "paragraph":
            # Add a paragraph with the specified font size
            paragraph = document.add_paragraph(text)
            # Set the font size for each run in the paragraph
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                
    # Save the document with the specified filename
    document.save(file_name)
    print(f"Document saved as {file_name}")


def generate(
    prompt: list,
    max_output_tokens: int = 2048,
    temperature: int = 2,
    top_p: float = 0.4,
    stream: bool = False,
) -> GenerationResponse | Iterable[GenerationResponse]:
    """
    Function to generate response using Gemini 1.5 Pro

    Args:
        prompt:
            List of prompt parts
        max_output_tokens:
            Max Output tokens
        temperature:
            Temperature for the model
        top_p
            Top-p for the model
        stream:
            Stream results?

    Returns:
        Model response """
    while True:
        try:
            responses = session.send_message(
                prompt,
                generation_config={
                    "max_output_tokens": max_output_tokens,
                    "temperature": temperature,
                    "top_p": top_p,
                },
                safety_settings={
                    HarmCategory.HARM_CATEGORY_HATE_SPEECH: BLOCK_LEVEL,
                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: BLOCK_LEVEL,
                    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: BLOCK_LEVEL,
                    HarmCategory.HARM_CATEGORY_HARASSMENT: BLOCK_LEVEL,
                },
                stream=stream,
            )
            break
        except Exception as e:
            print(f"An error occurred retrying in 2 seconds: {e}")
            time.sleep(2)

    return responses

def handle_step_one():
    print("Hello! I'm here to help build your bluesheet. What is the name of this project? Please upload the RFP document: ")
    project_name = input()
    rfp_document_path = input("What is the rfp document path: ")
    rfp_document = load_document(rfp_document_path)
    blue_sheet_template_path = input("What is the blue sheet template path: ")
    bluesheet_bid_doc_template = load_document(blue_sheet_template_path)
    
    prompt_template = f"""
<INSTRUCTIONS>
You are tasked with generating JSON data that represents the structure of a `.docx` document titled "<Project Name> – Basic RFP Bid Analysis". Replace <Project Name> with the name of the project (ensure it follows file naming conventions). This document will summarize essential project information as provided in the uploaded RFP. Each section should be formatted with clear headings and proper spacing for readability.
If user specifies the project name in <PROJECT_NAME> extract and use that name in the title. Otherwise, use the project name extracted from the RFP.

<FORMATTING GUIDELINES>
1. Title: Set the document title as "[Project Name] - RFP Analysis" using the project name extracted from the RFP.
2. Use the Blue_Sheet_Bid_Document_Template.pdf as a reference for formatting.
3. Use a header for each main section, with bold font for section names and sub-section names (e.g., "Bid Information:", "Funding Sources:").
4. Format content with appropriate spacing and alignment to ensure clarity.
5. Populate each field directly with extracted data from the RFP. If data is not available, label it as "Not specified".

<ACTION>
Parse and extract the following sections from the RFP document and format them as JSON according to this schema:

[{{"fileName": "<Project Name> - RFP Analysis.docx"}},
{{"type": "heading", "text": "<section title>", "fontSize": <font size>, "level": <heading level>}},
{{"type": "paragraph", "text": "<section content>", "fontSize": <font size>}},
...]

The sections are as follows:

1. Project Information:
   - Project Name
   - Location
   - Owner/Agency
   - Engineer or Consulting Firm (including contact information if available)

2. Bid Information:
   - Bid Number
   - Bid Date and Time
   - Pre-Bid Date, Time, and Type (mandatory or optional)

3. Funding Sources:
   - List the funding source(s) provided in the RFP (Federal, State, Local, Private, Turnkey, P3, or Other)

4. Project Scope:
   - Summarize the project's scope of work as described in the RFP, including major components or requirements.

5. Contractual Information:
   - Job Completion Date
   - Bid Acceptance Period
   - Guarantee Percentage
   - Construction Schedule (working days)
   - Estimated Award Date
   - Liquidated Damages (daily penalty cost)

6. Bid Conditions:
   - Listing Form requirement
   - Base Bid requirement
   - Substitutes Allowed (yes/no) with related comments
   - Or-Equals Accepted (yes/no) with related comments

<OUTPUT>
Make sure only valid JSON is generated with no errors.
Title: "<Project Name> - RFP Analysis"
Use the project name from the RFP as the title.

Territory:

Territory: [Determine the applicable territory based on the project location, such as Northern California, Southern California, etc. Match this territory with the sales or service regions defined by the organization].
Project Information:

Project Name: [Insert Project Name from RFP]
Location: [Insert project location, including city, county, and state]
Owner/Agency: [Insert name of the owner or agency managing the project]
Consulting / Engineer: [Insert name of the consulting or engineering firm overseeing the project]
Consulting Firm Contact: [Insert contact information if available, or any relevant details provided in the RFP]
Bid Information:

Bid Number: [Insert bid number if provided]
Bid Date and Time: [Insert the date and time for bid submission]
Pre-Bid Date: [Insert the date and details of any pre-bid meetings or conferences]
Funding Sources:
[List the funding sources as specified in the RFP, such as Federal, State, Local, Private, etc.]

Project Scope:
Scope of Work: [Summarize the project scope, including major components like infrastructure, construction, or system installation described in the RFP]

Contractual Information:

Job Completion Date: [Insert the number of working days or the expected completion timeline]
Bid Acceptance Period: [Insert the duration for which the bid is valid]
Guarantee Percentage: [Insert any guarantee or performance bond percentage required]
Liquidated Damages: [Insert the daily penalty or liquidated damages clause if stated]
Bid Conditions:

Listing Form Requirement: [Summarize subcontractor listing requirements, including thresholds and categories]
Base Bid Requirement: [Detail any base bid requirements, including forms and submission guidelines]
Substitutes Allowed: [Indicate whether material or product substitutions are allowed and under what conditions]
Or-Equals Accepted: [Indicate whether "or-equal" products or brands are accepted and describe the approval process]

<PROJECT_NAME>{project_name}</PROJECT_NAME>
"""
    # Request JSON content for the document
    responses1 = generate(prompt=[rfp_document, bluesheet_bid_doc_template, "Here is the rfp_document and bluesheet_bid_doc_template"])
    prompt = [prompt_template]
    print("Generating document structure in JSON format...")
    
    # Generate JSON response from prompt
    responses = generate(prompt=prompt)
    response_text = responses.text
    status = False
    while not status:
        if is_valid_json(response_text):
            try:
                # Save or pass the JSON data to the json_to_docx function

                json_to_docx(response_text)
                status = True
            except Exception as e: 
                print(f"An error occurred: {e}")
        else:
            print("Error occurred while generating the JSON structure. Trying again...")
            print(response_text)
            responses = generate(prompt=[
                """
                The document generation failed. Please try again, ensuring no errors in JSON structure.
                """
            ])
            response_text = responses.text
            if is_valid_json(response_text):
                try:
                    
                    json_to_docx(response_text)
                    status = True
                except Exception as e: 
                    print(f"An error occurred: {e}")

    

def handle_step_two():
    print("Would you like me to proceed with equipment identification for MISCO analysis?")
    response = input()
    misco_document = MISCO
    prompt_template = f"""
<INSTRUCTIONS>
You are tasked with performing an initial equipment identification and MISCO analysis. Using the RFP document provided by the user and MISCO Water Products, identify relevant equipment sections that match MISCO’s brands, product categories, and technologies. Your response should be clear, concise, and formatted for easy readability.

<FORMATTING GUIDELINES>
1. Begin with a statement confirming analysis completion, e.g., "Based on the provided RFP and MISCO’s product categories and brands, the following sections and equipment are relevant to MISCO:"
2. List each identified section and corresponding equipment description from the RFP, ensuring it aligns with MISCO’s represented brands and technologies.
3. Use bullet points or numbered lists for each item, with clear labels like "Spec Section:" and "Equipment Description:".
4. If no relevant sections are found, state "No relevant sections or equipment identified for MISCO."

<ACTION>
1. Review and cross-reference the RFP with 'MISCO Water Products' to identify sections and equipment relevant to MISCO.
2. Extract and list the identified spec sections and equipment descriptions.

<OUTPUT>
Example output format:
"Based on your provided RFP and MISCO's product categories and brands below, the following sections and equipment seem relevant to MISCO:"
• Spec Section: [specification number, e.g., "43 41 43"]
  Equipment Description: [brief description of equipment, e.g., "Fine Bubble Diffusers"]
• Spec Section: [another specification number]
  Equipment Description: [another equipment description]

<USER_RESPONSE>{response}</USER_RESPONSE>
If the user replies with "yes" or an affirmation, proceed with the extraction and analysis and display the output in the format above. If the user responds with "no" or a similar negative, respond with "Okay, I will not proceed with the analysis for MISCO."
"""

    prompt = [misco_document, prompt_template]
    responses = generate(prompt=prompt)
    response_text = responses.text
    print(response_text)



def handle_step_three():
    print("""
Would you like to identify synergy opportunities in this RFP for UFT’s platform companies,
Shape (representing related companies in Northern California),
and Southwest Valve? I will review the RFP based on their products and categories.
""")
    response = input()
    shape_document = SHAPE
    southwest_valve_document = SOUTH_WEST

    prompt_template = f"""
<INSTRUCTIONS>
You are tasked with identifying synergy opportunities. Analyze the RFP for potential opportunities for the following companies: 'Shape’s Represented Manufacturers and Categories' and 'Southwest Valve Manufacturers and Product Categories'  to identify relevant sections and equipment. Respond with a formatted summary of potential synergy opportunities, aligning each identified item with the specific companies and product categories.

<FORMATTING GUIDELINES>
1. Begin with a statement, e.g., "Based on the provided RFP and product categories from Shape and Southwest Valve, the following synergy opportunities are identified:"
2. Use bullet points for each synergy item, with labels like "Spec Section:" and "Relevant Equipment:".
3. Organize findings by company (Shape and Southwest Valve), listing relevant product categories and equipment in each section.
4. If no relevant synergies are found for either company, state "No relevant synergies identified for Shape or Southwest Valve."

<ACTION>
1. Analyze the RFP for references to equipment and technologies that match Shape’s and Southwest Valve's represented product categories.
2. For each identified match, list the relevant sections and equipment descriptions.
3. Structure output to clearly show synergy opportunities for Shape and Southwest Valve, based on the client's specifications.

<OUTPUT>
Example output format:
"Here are potential synergy opportunities for UFT's sister companies, Shape, and Southwest Valve based on the RFP: [List of identified sections and relevant equipment from UFT’s sister companies, Shape, and Southwest Valve]."

**Shape**
• Spec Section: [specification number, e.g., "45 12 30"]
  Relevant Equipment: [equipment name and description, e.g., "Variable Frequency Drives by ABB"]

• Spec Section: [another specification number]
  Relevant Equipment: [another equipment description]

**Southwest Valve**
• Spec Section: [specification number, e.g., "41 22 14"]
  Relevant Equipment: [equipment name and description, e.g., "Butterfly Valves (Metal-Seated)"]

<USER_RESPONSE>{response}</USER_RESPONSE>
If the user replies with "yes" or an affirmation, proceed with the analysis for synergy opportunities, using the RFP, Shape, and Southwest Valve product categories. Display the results as shown in the format above. If the user replies with "no" or a similar response, confirm by saying "Okay, I will not proceed with the synergy analysis for Shape and Southwest Valve."
"""
    prompt = [shape_document, southwest_valve_document, prompt_template]
    responses = generate(prompt=prompt)
    response_text = responses.text
    print(response_text)


def handle_step_four():
    bluesheet_template = load_document("chat-agent\Blue Sheet Template 2024.csv")
    print("I will go ahead to generate a draft of the bluesheet in Excel format.")
   
    # Initial prompt for generating JSON for the bluesheet draft
    prompt_template_one = """
<INSTRUCTIONS>
Generate JSON output to create an Excel document draft based on the structure of "Blue Sheet Template 2024.csv". This JSON should represent the data for three separate worksheets—one each for MISCO, Shape, and Southwest Valve. Each worksheet should be structured as a list of dictionaries, where each dictionary represents a row with key-value pairs for column headers and cell values.

<JSON SCHEMA>
Your JSON output should have the following structure:
{
    "document_name": "<Project Name> – Bluesheet Draft.xlsx",
    "MISCO": [
        { "Spec Section": "...", "Equipment/Item Description": "...", "Named Manufacturers": "...", "Represented Company": "...", "Contact Information": "...", "Product Specifications": "..." },
        ...
    ],
    "Shape": [
        { "Spec Section": "...", "Equipment/Item Description": "...", "Named Manufacturers": "...", "Represented Company": "...", "Contact Information": "...", "Product Specifications": "..." },
        ...
    ],
    "Southwest Valve": [
        { "Spec Section": "...", "Equipment/Item Description": "...", "Named Manufacturers": "...", "Represented Company": "...", "Contact Information": "...", "Product Specifications": "..." },
        ...
    ]
}

<FORMATTING GUIDELINES>
1. Provide three worksheets with headers (e.g., Spec Section, Equipment/Item Description, Named Manufacturers, Represented Company, Contact Information, Product Specifications).
2. Each row should be formatted as a dictionary entry with the appropriate data extracted from the RFP.
3. For each worksheet (MISCO, Shape, Southwest Valve):
   - Include fields for spec sections, equipment descriptions, manufacturers, represented companies, and contacts.
   - Extract data based on relevance for each company, and if data is missing, label as "Not specified".
   - Example row for JSON entry:
      { "Spec Section": "33 40 00", "Equipment/Item Description": "Valve", "Named Manufacturers": "Brand A, Brand B", "Represented Company": "MISCO", "Contact Information": "Sandy Clarke, sclarke@miscowater.com", "Product Specifications": "Max pressure: 100 PSI" }

<OUTPUT>
Generate only the JSON formatted as per the schema above. Save it with the specified structure to represent each company's sheet data clearly.
"""

    # Request initial JSON for generating the bluesheet draft
    prompt1 = [bluesheet_template, prompt_template_one]
    responses = generate(prompt=prompt1)
    response_text = responses.text

    # Check if response is JSON and attempt to save it
    status = False
    if is_valid_json(response_text):
        bluesheet_data = json.loads(re.sub(r'```json|```', '', response_text).strip())
        try:
            save_to_excel(bluesheet_data)
            status = True
        except Exception as e:
            print(f"Error saving bluesheet: {e}")

    # Loop for handling user feedback and modifications
    while True:
        if status:
            # Success case - ask for user review
            responses = generate(prompt=[
                """
                <STATUS>"Success"</STATUS>
                <ACTION>Ask user to review draft to see if there are modifications
                <MESSAGE FOR USER> "I have generated a JSON draft for the bluesheet, ready to be converted to Excel format with detailed product specifications, criteria, and requirements for each category, as well as the contact information for MISCO, Shape, and Southwest Valves opportunities. Please review this draft and let me know if there are any changes or additions."
                """
            ])
        else:
            # Failure case - regenerate JSON without errors
            responses = generate(prompt=[
                """
                <STATUS>"Fail"</STATUS>
                <ACTION>Regenerate JSON output, following the correct schema for the Excel document generation.
                """
            ])

        response_text = responses.text
        if is_valid_json(response_text):
            bluesheet_data = json.loads(re.sub(r'```json|```', '', response_text).strip())
            try:
                save_to_excel(bluesheet_data)
                status = True
                continue
            except Exception as e:
                print(f"Error saving bluesheet: {e}")
                continue
        else:
            print(response_text)
            response = input("Please provide your modifications: ")
            responses = generate(prompt=[
                f"""
                <INSTRUCTIONS>
                Apply user feedback to the JSON draft, updating or adding rows to the specified worksheet (MISCO, Shape, or Southwest Valve) in the schema.
                <USER_RESPONSE>{response}</USER_RESPONSE>
                """
            ])
            response_text = responses.text
            status = False
            if is_valid_json(response_text):
                bluesheet_data = json.loads(re.sub(r'```json|```', '', response_text).strip())
                try:
                    save_to_excel(bluesheet_data)
                    status = True
                except Exception as e:
                    print(f"Error saving bluesheet: {e}")
                continue
            else:
                print(response_text)
                break

def handle_step_five():
    # Initial prompt to ask the user if they would like an email draft
    user_response = input("Would you like to draft an email to UFT portfolio company leads? (yes/no): ").strip().lower()

    # Prepare the appropriate prompt based on the user's response
    if user_response == "yes":
        # Prompt template for drafting the initial email
        email_draft_prompt = f"""
<INSTRUCTIONS>
Draft an email addressed to UFT portfolio company leads based on the bluesheet document generated in previous steps. This email should summarize the analysis conducted and offer the bluesheet as an attachment for further review. Ensure the email is professional and contains key details for the recipient to understand the document's contents.

<FORMATTING GUIDELINES>
1. Use a clear and professional subject line, e.g., "RFP Analysis and Bluesheet for [Project Name]".
2. Start with a greeting, addressing UFT portfolio company leads directly (e.g., "Dear UFT Portfolio Company Leads").
3. Summarize the purpose of the email in a few sentences, highlighting that the email contains the bluesheet for the RFP analysis.
4. Briefly describe the contents of the bluesheet, mentioning key sections like project information, bid information, funding sources, equipment details, and synergy opportunities with MISCO, Shape, and Southwest Valve.
5. Conclude with a closing statement offering further assistance or requesting feedback if needed, followed by a professional signature (e.g., "Best regards, [Your Name]").

<ACTION>
Generate a text-based email draft for user review, formatted according to the guidelines.

<USER_RESPONSE>
{user_response}
</USER_RESPONSE>
"""

        # Generate the email draft based on the user input
        email_draft = generate(prompt=[email_draft_prompt])
        email_text = email_draft.text
        print("Draft email generated:")
        print(email_text)  # Display email draft to user

        # Ask the user if they would like any modifications to the draft
        user_response = input("Would you like any modifications to the email draft? (yes/no): ").strip().lower()

        while user_response == "yes":
            # Get specific modifications from the user
            modification_details = input("Please specify the modifications you'd like to make: ")

            # Prompt template for modifications to the email draft
            email_modification_prompt = f"""
<INSTRUCTIONS>
Make the following modifications to the email draft based on user feedback. Keep the structure and formatting professional and clear. Implement any specific changes requested by the user, ensuring that the email retains its original purpose and flow.

<USER_RESPONSE>
{modification_details}
</USER_RESPONSE>

<FORMATTING GUIDELINES>
Maintain the initial formatting and organization of the email, ensuring all modifications are consistent with a professional business email.

<ACTION>
Apply user-specified changes to the draft email, revising sections as necessary. Generate an updated text-based draft incorporating these changes.

<OUTPUT>
The modified email draft reflecting the requested changes.
"""

            # Generate the modified email draft
            modified_email_draft = generate(prompt=[email_modification_prompt])
            modified_email_text = modified_email_draft.text
            print("Modified draft email:")
            print(modified_email_text)  # Display modified email draft to user

            # Ask if further modifications are needed
            user_response = input("Would you like any further modifications to the email draft? (yes/no): ").strip().lower()

        # Finalize the email draft if no more modifications are needed
        print("The email draft is finalized and ready for sending.")
        return modified_email_text if user_response == "yes" else email_text  # Return the finalized email draft for sending or saving

    elif user_response == "no":
        # If the user does not want an email draft, prepare the final bluesheet in JSON format
        bluesheet_finalization_prompt = f"""
<INSTRUCTIONS>
Prepare the final bluesheet in JSON format based on the bluesheet draft generated in previous steps. The output should follow this JSON schema:

{
    "excel_document_name": "[Project Name] – Final Bluesheet.xlsx",
    "MISCO": [{ "column1": "value", "column2": "value", ... }],
    "Shape": [{ "column1": "value", "column2": "value", ... }],
    "Southwest Valve": [{ "column1": "value", "column2": "value", ... }]
}

Ensure the JSON is formatted with lists of dictionaries for each company (MISCO, Shape, Southwest Valve), where each dictionary represents a row.

<ACTION>
Generate JSON following the schema specified. Use re.sub(r'```json|```', '', response_text).strip() to extract JSON output from the response.
"""

        # Generate the final bluesheet JSON
        bluesheet_json_response = generate(prompt=[bluesheet_finalization_prompt])
        bluesheet_json_text = re.sub(r'```json|```', '', bluesheet_json_response.text).strip()

        # Save the JSON output as an Excel file
        try:
            bluesheet_data = json.loads(bluesheet_json_text)
            # Save Excel file using openpyxl based on the JSON schema
            book_name = bluesheet_data["excel_document_name"]
            workbook = openpyxl.Workbook()
            for company in ["MISCO", "Shape", "Southwest Valve"]:
                if company in bluesheet_data:
                    worksheet = workbook.create_sheet(company)
                    # Write data to worksheet
                    for row_data in bluesheet_data[company]:
                        worksheet.append([row_data[col] for col in row_data])
            workbook.save(f"{book_name}.xlsx")
            print("The final bluesheet Excel file has been successfully generated and saved.")
        except Exception as e:
            print(f"Error generating the final bluesheet: {e}")
    else:
        # Handle unexpected responses
        print("Invalid response. Please respond with 'yes' to draft an email or 'no' to proceed with bluesheet download.")


def handle_step_six():
    while True:
        # Initial prompt to ask if the user would like to do anything else
        follow_up_prompt = """
    <INSTRUCTIONS>Ask user if they would like to perform any other actions or need further assistance.
    <Question>
    Would you like to perform any other actions or need further assistance?
    <Action if Yes>
    Provide details on the next action you would like to take, and I will assist you.
    <Action if No>
    I will conclude the session. Thank you for using this service.
    <Output>
    If user responds with "Yes", proceed to ask for specific details on the requested action. If "No", confirm session completion.

    """

        # Send the prompt to the model to capture the user's intent (yes/no response)
        follow_up_response = generate(prompt=[follow_up_prompt])
        response_text = follow_up_response.text
        print(response_text)  # Display model's follow-up prompt to user

        # Capture the user's actual response to the follow-up prompt
        user_response = input().strip().lower()

        # Pass the user's response back to the model for interpretation
        response_handling_prompt = f"""
    <INSTRUCTIONS>
    The user has responded to the question about whether they need further assistance. Interpret the response, asking for specifics if they would like additional actions or concluding the session if they do not.
    <USER_RESPONSE>
    {user_response}
    <Output>
    If the response is affirmative, ask for details on the next action. If the response is negative, provide a thank-you message to end the session. Start the thank-you message with "END: Thank you for using this service."
    """

        # Generate the model's handling of the user's response
        final_follow_up = generate(prompt=[response_handling_prompt])
        print(final_follow_up.text)  # Display model's final follow-up or conclusion message

        if final_follow_up.text.startswith("END"):
            break




while True:
    print("Welcome to the Bluesheet Generator! Let's get started.(yes/no)")
    if input().strip().lower() == "no":
        print("Thank you for using the Bluesheet Generator. Goodbye!")
        break
    model = GenerativeModel(
    MODEL_NAME,
    system_instruction=SYSTEM_INSTRUCTION
    )

    session = model.start_chat()
    handle_step_one()
    handle_step_two()
    handle_step_three()
    handle_step_four()
    handle_step_five()
    handle_step_six()
